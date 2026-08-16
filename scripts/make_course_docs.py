"""Build the E2E course material.

Equations are written as real OMML (m:oMath), not as plain text, because that is
the path the maths-aware ingestion actually has to handle -- python-docx's
paragraph.text silently drops m:oMath, which is the defect the omml.py converter
exists to fix. Plain-text formulas would pass the test without exercising it.
"""
import os
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

nsmap.setdefault("m", "http://schemas.openxmlformats.org/officeDocument/2006/math")

# Writes straight into the frontend's Playwright fixtures. Override with
# E2E_FIXTURES when the frontend checkout is not a sibling of this one.
OUT = Path(
    os.getenv("E2E_FIXTURES")
    or Path(__file__).resolve().parents[2] / "merakiai-frontend" / "e2e" / "fixtures"
)
OUT.mkdir(parents=True, exist_ok=True)


def m(tag):
    return OxmlElement(f"m:{tag}")


def run(text):
    r = m("r")
    t = m("t")
    t.text = text
    # keep spaces the converter relies on for token separation
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def frac(num_children, den_children):
    f = m("f")
    num, den = m("num"), m("den")
    for c in num_children:
        num.append(c)
    for c in den_children:
        den.append(c)
    f.append(num)
    f.append(den)
    return f


def sup(base_children, exp_children):
    s = m("sSup")
    e, sup_ = m("e"), m("sup")
    for c in base_children:
        e.append(c)
    for c in exp_children:
        sup_.append(c)
    s.append(e)
    s.append(sup_)
    return s


def sub(base_children, sub_children):
    s = m("sSub")
    e, sub_ = m("e"), m("sub")
    for c in base_children:
        e.append(c)
    for c in sub_children:
        sub_.append(c)
    s.append(e)
    s.append(sub_)
    return s


def nary(op, lower, upper, body):
    """n-ary operator (integral, sum) with limits."""
    n = m("nary")
    pr = m("naryPr")
    chr_ = m("chr")
    chr_.set(qn("m:val"), op)
    pr.append(chr_)
    limloc = m("limLoc")
    limloc.set(qn("m:val"), "subSup")
    pr.append(limloc)
    n.append(pr)
    for tag, children in (("sub", lower), ("sup", upper), ("e", body)):
        el = m(tag)
        for c in children:
            el.append(c)
        n.append(el)
    return n


def lim(var, target, body):
    """limit as a lower-limit function."""
    fn = m("func")
    fname = m("fName")
    limlow = m("limLow")
    e, low = m("e"), m("lim")
    e.append(run("lim"))
    low.append(run(f"{var}\u2192{target}"))
    limlow.append(e)
    limlow.append(low)
    fname.append(limlow)
    fn.append(fname)
    body_el = m("e")
    for c in body:
        body_el.append(c)
    fn.append(body_el)
    return fn


def add_math(doc, children, display=True):
    """Append a paragraph holding one OMML expression."""
    p = doc.add_paragraph()
    if display:
        omath_para = m("oMathPara")
        omath = m("oMath")
        for c in children:
            omath.append(c)
        omath_para.append(omath)
        p._p.append(omath_para)
    else:
        omath = m("oMath")
        for c in children:
            omath.append(c)
        p._p.append(omath)
    return p


def add_inline(doc, before, children, after):
    """A sentence with an inline equation inside it."""
    p = doc.add_paragraph()
    p.add_run(before)
    omath = m("oMath")
    for c in children:
        omath.append(c)
    p._p.append(omath)
    p.add_run(after)
    return p


# ─── Document 1: Learn — lecture notes ───────────────────────────────────────
doc = Document()
doc.add_heading("Calculus I — Differentiation", level=1)
doc.add_paragraph(
    "These notes cover the derivative as a limit, the standard differentiation "
    "rules, and the chain rule. They are written for Level 100 students taking "
    "Calculus I at the University of Ghana."
)

doc.add_heading("1. The derivative as a limit", level=2)
doc.add_paragraph(
    "The derivative of a function measures its instantaneous rate of change. "
    "Formally, the derivative of f at the point x is defined by the limit of the "
    "difference quotient as the increment h approaches zero:"
)
add_math(doc, [
    run("f'(x) = "),
    lim("h", "0", [frac([run("f(x+h) - f(x)")], [run("h")])]),
])
doc.add_paragraph(
    "If this limit exists, f is said to be differentiable at x. A function that "
    "is differentiable at a point is necessarily continuous there, but the "
    "converse is false: the absolute value function is continuous at zero and "
    "has no derivative there, because the left and right limits of the "
    "difference quotient disagree."
)

doc.add_heading("2. The power rule", level=2)
doc.add_paragraph("For any real exponent n, the derivative of a power of x is:")
add_math(doc, [
    frac([run("d")], [run("dx")]),
    sup([run("x")], [run("n")]),
    run(" = n"),
    sup([run("x")], [run("n-1")]),
])
add_inline(doc,
           "So the derivative of ",
           [sup([run("x")], [run("3")])],
           " is 3x², and the derivative of the square root of x, written as a "
           "power, follows the same rule.")
add_math(doc, [
    frac([run("d")], [run("dx")]),
    sup([run("x")], [frac([run("1")], [run("2")])]),
    run(" = "),
    frac([run("1")], [run("2"), sup([run("x")], [frac([run("1")], [run("2")])])]),
])

doc.add_heading("3. Product and quotient rules", level=2)
doc.add_paragraph(
    "The derivative of a product is not the product of the derivatives. This is "
    "the single most common error in first-year scripts. The correct rule is:"
)
add_math(doc, [run("(uv)' = u'v + uv'")])
doc.add_paragraph("For a quotient, the rule is:")
add_math(doc, [
    run("("),
    frac([run("u")], [run("v")]),
    run(")' = "),
    frac([run("u'v - uv'")], [sup([run("v")], [run("2")])]),
])

doc.add_heading("4. The chain rule", level=2)
doc.add_paragraph(
    "The chain rule differentiates a composition of functions. If y depends on u "
    "and u depends on x, then the rate of change of y with respect to x is the "
    "product of the two rates:"
)
add_math(doc, [
    frac([run("dy")], [run("dx")]),
    run(" = "),
    frac([run("dy")], [run("du")]),
    run(" \u00b7 "),
    frac([run("du")], [run("dx")]),
])
doc.add_paragraph("Equivalently, in function notation:")
add_math(doc, [run("[f(g(x))]' = f'(g(x)) \u00b7 g'(x)")])

doc.add_heading("Worked example 1", level=3)
doc.add_paragraph("Differentiate y = (3x² + 1)⁵.")
doc.add_paragraph("Step 1. Identify the outer and inner functions. Let u = 3x² + 1, so y = u⁵.")
doc.add_paragraph("Step 2. Differentiate each part separately.")
add_math(doc, [
    frac([run("dy")], [run("du")]),
    run(" = 5"),
    sup([run("u")], [run("4")]),
    run(",    "),
    frac([run("du")], [run("dx")]),
    run(" = 6x"),
])
doc.add_paragraph("Step 3. Multiply and substitute u back.")
add_math(doc, [
    frac([run("dy")], [run("dx")]),
    run(" = 5"),
    sup([run("(3"), sup([run("x")], [run("2")]), run(" + 1)")], [run("4")]),
    run(" \u00b7 6x = 30x"),
    sup([run("(3"), sup([run("x")], [run("2")]), run(" + 1)")], [run("4")]),
])

doc.add_heading("5. Standard limits", level=2)
doc.add_paragraph("Two limits are used repeatedly when differentiating trigonometric functions:")
add_math(doc, [
    lim("x", "0", [frac([run("sin x")], [run("x")])]),
    run(" = 1"),
])
add_math(doc, [
    lim("x", "0", [frac([run("1 - cos x")], [run("x")])]),
    run(" = 0"),
])

doc.add_heading("6. Table of standard derivatives", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Function"
hdr[1].text = "Derivative"
for fn, dv in [
    ("sin x", "cos x"),
    ("cos x", "-sin x"),
    ("tan x", "sec² x"),
    ("eˣ", "eˣ"),
    ("ln x", "1/x"),
]:
    row = table.add_row().cells
    row[0].text = fn
    row[1].text = dv

doc.save(OUT / "Calculus-I-Differentiation-Notes.docx")
print("wrote", OUT / "Calculus-I-Differentiation-Notes.docx")


# ─── Document 2: Review — tutorial questions ─────────────────────────────────
doc2 = Document()
doc2.add_heading("Calculus I — Tutorial Questions and Mark Scheme", level=1)
doc2.add_paragraph(
    "Tutorial sheet 3. Attempt all questions. Marks are shown in brackets. This "
    "sheet covers differentiation from first principles, the product and "
    "quotient rules, and the chain rule."
)

doc2.add_heading("Section A — Short questions", level=2)
doc2.add_paragraph("A1. Differentiate from first principles. [5 marks]", style="List Number")
add_inline(doc2, "Find the derivative of ",
           [sup([run("f(x) = x")], [run("2")])],
           " using the limit definition of the derivative.")
doc2.add_paragraph(
    "Mark scheme: correct difference quotient (2), expansion of (x+h)² (1), "
    "cancellation of h (1), limit evaluated to 2x (1)."
)

doc2.add_paragraph("A2. Product rule. [4 marks]", style="List Number")
add_inline(doc2, "Differentiate ", [run("y = x")], "² sin x with respect to x.")
doc2.add_paragraph("Mark scheme: identifies u and v (1), applies (uv)' = u'v + uv' (2), final answer 2x sin x + x² cos x (1).")

doc2.add_paragraph("A3. Quotient rule. [4 marks]", style="List Number")
add_math(doc2, [
    run("y = "),
    frac([run("sin x")], [run("x")]),
])
doc2.add_paragraph("Mark scheme: correct quotient rule statement (2), correct substitution (1), simplified answer (1).")

doc2.add_heading("Section B — Chain rule", level=2)
doc2.add_paragraph("B1. Differentiate the following. [3 marks each]", style="List Number")
add_math(doc2, [run("y = "), sup([run("(2x - 5)")], [run("7")])])
add_math(doc2, [run("y = sin(3"), sup([run("x")], [run("2")]), run(")")])
add_math(doc2, [run("y = "), sup([run("e")], [run("4x")])])
doc2.add_paragraph(
    "Mark scheme: identifies inner function (1), differentiates outer correctly (1), "
    "multiplies by the derivative of the inner function (1). A common error is "
    "omitting the inner derivative entirely — award zero for the final mark."
)

doc2.add_heading("Section C — Applications", level=2)
doc2.add_paragraph(
    "C1. A spherical balloon is inflated so that its radius increases at 2 cm/s. "
    "Find the rate at which the volume increases when the radius is 5 cm. [6 marks]"
)
add_math(doc2, [
    run("V = "),
    frac([run("4")], [run("3")]),
    run("\u03c0"),
    sup([run("r")], [run("3")]),
    run(",    "),
    frac([run("dV")], [run("dt")]),
    run(" = 4\u03c0"),
    sup([run("r")], [run("2")]),
    frac([run("dr")], [run("dt")]),
])
doc2.add_paragraph("Mark scheme: correct volume formula (1), differentiates with respect to t using the chain rule (3), substitutes r = 5 and dr/dt = 2 (1), answer 200π cm³/s (1).")

doc2.save(OUT / "Calculus-I-Tutorial-Questions.docx")
print("wrote", OUT / "Calculus-I-Tutorial-Questions.docx")


# ─── Document 3: Application — case studies ──────────────────────────────────
doc3 = Document()
doc3.add_heading("Calculus I — Real-World Applications", level=1)
doc3.add_paragraph(
    "Case studies connecting differentiation to problems students meet in "
    "engineering, economics and the sciences in Ghana."
)

doc3.add_heading("Case study 1: Optimising a cocoa storage shed", level=2)
doc3.add_paragraph(
    "A cooperative in the Ashanti Region needs a rectangular storage shed with a "
    "fixed floor area of 200 square metres. Walling costs GHS 90 per metre of "
    "perimeter. The cooperative wants the cheapest shed. Because the area is "
    "fixed, the width depends on the length, and the perimeter becomes a "
    "function of one variable:"
)
add_math(doc3, [
    run("P(x) = 2x + "),
    frac([run("400")], [run("x")]),
])
doc3.add_paragraph(
    "Setting the derivative to zero locates the minimum. This is a genuine "
    "optimisation: the derivative is not merely a formula to be memorised but "
    "the tool that finds the cheapest design."
)
add_math(doc3, [
    run("P'(x) = 2 - "),
    frac([run("400")], [sup([run("x")], [run("2")])]),
    run(" = 0"),
])

doc3.add_heading("Case study 2: Marginal cost in a bakery", level=2)
doc3.add_paragraph(
    "In economics the derivative of a total cost function is the marginal cost — "
    "the cost of producing one more unit. A bakery in Accra models its daily "
    "cost of producing q loaves as a quadratic function. The marginal cost is "
    "the derivative, and the point where marginal cost equals price is where "
    "profit stops increasing."
)
add_math(doc3, [
    run("C(q) = 500 + 3q + 0.01"),
    sup([run("q")], [run("2")]),
    run(",    MC(q) = C'(q) = 3 + 0.02q"),
])

doc3.add_heading("Case study 3: Drug concentration and rates of change", level=2)
doc3.add_paragraph(
    "The concentration of a drug in the bloodstream falls exponentially after a "
    "dose. The derivative gives the rate at which the drug is cleared, which "
    "determines how often a dose must be repeated."
)
add_math(doc3, [
    run("C(t) = "),
    sub([run("C")], [run("0")]),
    sup([run("e")], [run("-kt")]),
    run(",    C'(t) = -k"),
    sub([run("C")], [run("0")]),
    sup([run("e")], [run("-kt")]),
])

doc3.add_heading("Case study 4: Velocity and acceleration of a vehicle", level=2)
doc3.add_paragraph(
    "For a vehicle on the Accra–Kumasi road whose displacement is given as a "
    "function of time, the first derivative is velocity and the second "
    "derivative is acceleration. Where the velocity is zero and changes sign, "
    "the vehicle reverses direction."
)
add_math(doc3, [
    run("v(t) = "),
    frac([run("ds")], [run("dt")]),
    run(",    a(t) = "),
    frac([run("dv")], [run("dt")]),
    run(" = "),
    frac([run("d"), sup([run("")], [run("2")]), run("s")], [run("d"), sup([run("t")], [run("2")])]),
])

doc3.save(OUT / "Calculus-I-Applications.docx")
print("wrote", OUT / "Calculus-I-Applications.docx")
