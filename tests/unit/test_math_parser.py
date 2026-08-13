"""Maths-preserving document parsing.

The regression these guard against: the previous pipeline read a lecturer's
Word notes, dropped every equation, and joined whole sections into a single
un-splittable string. A worked example arrived at the model as "Example 1:
Evaluate the limit below." / "Step 1: Factor the numerator." with the
expression itself missing.
"""

import io

import docx
import pytest
from docx.oxml import parse_xml

from app.ai.ingestion.math_parser import (
    EQUATION,
    HEADING,
    LIST_ITEM,
    PARAGRAPH,
    TABLE,
    parse_blocks,
)

NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'


def omath(inner: str):
    return parse_xml(f"<m:oMath {NS}>{inner}</m:oMath>")


def mrun(text: str) -> str:
    return f"<m:r><m:t>{text}</m:t></m:r>"


def power(base: str, exponent: str) -> str:
    return f"<m:sSup><m:e>{mrun(base)}</m:e><m:sup>{mrun(exponent)}</m:sup></m:sSup>"


def to_bytes(document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def notes():
    """A miniature version of real Calculus lecture notes."""
    d = docx.Document()
    d.add_heading("Differentiation", level=1)
    d.add_heading("The Power Rule", level=2)

    p = d.add_paragraph()
    p.add_run("The derivative of ")
    p._p.append(omath(power("x", "n")))
    p.add_run(" is ")
    p._p.append(omath(mrun("n") + power("x", "n-1")))
    p.add_run(".")

    bullet = d.add_paragraph(style="List Bullet")
    bullet.add_run("Valid whenever ")
    bullet._p.append(omath(mrun("n") + mrun("≠") + mrun("0")))

    standalone = d.add_paragraph()
    standalone._p.append(omath(power("e", "x")))

    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Function"
    table.rows[0].cells[1].text = "Derivative"
    table.rows[1].cells[0].text = "sin x"
    table.rows[1].cells[1].text = "cos x"

    return parse_blocks(to_bytes(d), "calculus.docx")


def test_equations_are_not_dropped(notes):
    latex = " ".join(b["text"] for b in notes)
    assert "x^n" in latex
    assert "nx^{n-1}" in latex


def test_inline_maths_keeps_document_order(notes):
    """Maths must stay where it was written, not be appended elsewhere."""
    para = next(b for b in notes if b["type"] == PARAGRAPH and "derivative" in b["text"])
    assert para["text"] == "The derivative of $x^n$ is $nx^{n-1}$."


def test_heading_hierarchy_is_recorded(notes):
    para = next(b for b in notes if b["type"] == PARAGRAPH and "derivative" in b["text"])
    assert para["heading_path"] == ["Differentiation", "The Power Rule"]


def test_headings_become_their_own_blocks(notes):
    headings = [b["text"] for b in notes if b["type"] == HEADING]
    assert headings == ["Differentiation", "The Power Rule"]


def test_equation_only_paragraph_becomes_display_maths(notes):
    block = next(b for b in notes if b["type"] == EQUATION)
    assert block["text"] == "$$\ne^x\n$$"
    assert block["has_math"] is True


def test_list_items_are_detected_from_style(notes):
    bullet = next(b for b in notes if b["type"] == LIST_ITEM)
    assert r"n\neq 0" in bullet["text"]


def test_table_structure_is_preserved_as_markdown(notes):
    table = next(b for b in notes if b["type"] == TABLE)
    assert "| Function | Derivative |" in table["text"]
    assert "| sin x | cos x |" in table["text"]


def test_table_cells_convert_their_equations():
    d = docx.Document()
    table = d.add_table(rows=1, cols=1)
    table.rows[0].cells[0].paragraphs[0]._p.append(omath(power("e", "x")))
    blocks = parse_blocks(to_bytes(d), "t.docx")
    block = next(b for b in blocks if b["type"] == TABLE)
    assert "$e^x$" in block["text"]
    assert block["has_math"] is True


def test_short_formula_section_is_kept():
    """The old 40-word floor silently deleted formula-dense material."""
    d = docx.Document()
    d.add_heading("Key Identity", level=1)
    p = d.add_paragraph()
    p._p.append(omath(power("a", "2")))
    blocks = parse_blocks(to_bytes(d), "short.docx")
    assert any(b["has_math"] for b in blocks)


def test_literal_dollar_sign_is_escaped():
    d = docx.Document()
    d.add_paragraph("The bond costs $500 at maturity.")
    blocks = parse_blocks(to_bytes(d), "money.docx")
    assert r"\$500" in blocks[0]["text"]


def test_blocks_are_indexed_in_document_order(notes):
    assert [b["index"] for b in notes] == sorted(b["index"] for b in notes)


def test_markdown_source_keeps_existing_latex():
    source = "# Integrals\n\nWe know $\\int_0^1 x^2 dx = \\frac{1}{3}$ exactly.\n"
    blocks = parse_blocks(source.encode(), "notes.md")
    para = next(b for b in blocks if b["type"] == PARAGRAPH)
    assert para["has_math"] is True
    assert para["heading_path"] == ["Integrals"]


def test_unsupported_extension_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        parse_blocks(b"data", "notes.pptx")


class TestPdfBoilerplate:
    """Running headers/footers must not be mistaken for section headings.

    A footer like "1 | P a g e" repeats on every page and, being short and
    title-ish, satisfies the heading heuristic — which then attaches it as the
    heading path of every following paragraph.
    """

    @staticmethod
    def _page(n: int) -> str:
        return (
            f"{n} | P a g e\n"
            f"Opening sentence of page {n}.\n"
            f"A middle sentence carrying the argument forward.\n"
            f"A closing sentence for page {n}.\n"
        )

    def test_repeated_footer_is_detected(self):
        from app.ai.ingestion.math_parser import _find_boilerplate

        boilerplate = _find_boilerplate([self._page(n) for n in range(1, 6)])
        assert _fingerprint("3 | P a g e") in boilerplate

    def test_body_text_is_not_treated_as_boilerplate(self):
        """Templated body lines differing only by digits must survive."""
        from app.ai.ingestion.math_parser import _find_boilerplate

        boilerplate = _find_boilerplate([self._page(n) for n in range(1, 6)])
        assert _fingerprint("Opening sentence of page 2.") not in boilerplate

    def test_sparse_pages_contribute_no_candidates(self):
        """With only a couple of lines on a page, everything is content."""
        from app.ai.ingestion.math_parser import _find_boilerplate

        pages = [f"Chapter {n}\nA paragraph about topic {n}." for n in range(1, 6)]
        assert _find_boilerplate(pages) == set()

    def test_single_page_document_has_no_boilerplate(self):
        from app.ai.ingestion.math_parser import _find_boilerplate

        assert _find_boilerplate(["1 | P a g e\nOnly page."]) == set()


def _fingerprint(line: str) -> str:
    from app.ai.ingestion.math_parser import _boilerplate_fingerprint

    return _boilerplate_fingerprint(line)
