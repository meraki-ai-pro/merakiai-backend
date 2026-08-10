from docx import Document
import tempfile
import os


def _load_partition():
    """Import unstructured's auto-partitioner lazily.

    ``unstructured.file_utils.filetype`` does ``importlib.import_module("magic")``
    at module scope. On Windows, python-magic hunts for libmagic DLLs and can
    spin at 100% CPU indefinitely when it cannot find them — which blocked
    ``app.main`` from ever finishing its import, so the API never bound a port.
    Deferring the import to the one function that needs it keeps startup at
    ~11s and confines any such hang to PDF parsing inside the ingestion worker.
    """
    from unstructured.partition.auto import partition

    return partition


def parse_document(upload_file):
    """
    Entry point for parsing documents.
    Supports PDF and DOCX.
    Returns structured sections.
    """

    suffix = upload_file.filename.split(".")[-1].lower()

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
        tmp.write(upload_file.file.read())
        tmp_path = tmp.name

    try:
        if suffix == "pdf":
            return parse_pdf(tmp_path)

        elif suffix in ("docx", "doc"):
            return parse_docx(tmp_path)

        else:
            raise ValueError("Unsupported file type")

    finally:
        os.remove(tmp_path)

def parse_pdf(path):
    """
    Uses unstructured for layout-aware parsing.
    Falls back to PyPDF if needed.
    """
    elements = _load_partition()(filename=path)

    sections = []
    current_section = {"title": None, "content": ""}

    for el in elements:
        text = el.text.strip() if el.text else ""
        if not text:
            continue

        # Detect headings
        if el.category in ("Title", "Header"):
            if current_section["content"]:
                sections.append(current_section)
            current_section = {
                "title": text,
                "content": ""
            }
        else:
            current_section["content"] += text + "\n"

    if current_section["content"]:
        sections.append(current_section)

    return sections

def parse_docx(path):
    doc = Document(path)

    sections = []
    current_section = {"title": None, "content": ""}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style.name.startswith("Heading"):
            if current_section["content"]:
                sections.append(current_section)
            current_section = {
                "title": text,
                "content": ""
            }
        else:
            current_section["content"] += text + "\n"

    # Handle tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            current_section["content"] += row_text + "\n"

    if current_section["content"]:
        sections.append(current_section)

    return sections
